"""
Land Contract Generator
========================
Fills [[PLACEHOLDERS]] in the original Word template and converts to PDF
using LibreOffice. Output is pixel-perfect match to the original document.

Template file: backend/doc_templates/TEMPLATE_land_contract.docx
LibreOffice must be installed on the server.
"""

import os
import shutil
import subprocess
import tempfile
from docx import Document


# ── Number to words ────────────────────────────────────────────────────────────
_ones = ['','One','Two','Three','Four','Five','Six','Seven','Eight','Nine',
         'Ten','Eleven','Twelve','Thirteen','Fourteen','Fifteen','Sixteen',
         'Seventeen','Eighteen','Nineteen']
_tens = ['','','Twenty','Thirty','Forty','Fifty','Sixty','Seventy','Eighty','Ninety']

def _n2w(n):
    n = int(n)
    if n == 0:  return 'Zero'
    if n < 20:  return _ones[n]
    if n < 100: return _tens[n//10] + ('-' + _ones[n%10] if n%10 else '')
    if n < 1_000:
        return _ones[n//100] + ' Hundred' + (' ' + _n2w(n%100) if n%100 else '')
    if n < 1_000_000:
        return _n2w(n//1000) + ' Thousand' + (' ' + _n2w(n%1000) if n%1000 else '')
    return _n2w(n//1_000_000) + ' Million' + (' ' + _n2w(n%1_000_000) if n%1_000_000 else '')

def d2w(amount):
    """95000.00 -> 'Ninety-Five Thousand and no/100'"""
    dollars = int(amount)
    cents   = round((amount - dollars) * 100)
    return _n2w(dollars) + (f' and {cents}/100' if cents else ' and no/100')

def n2w(n):
    return _n2w(int(n))

def pct2w(rate):
    whole = int(rate)
    frac  = rate - whole
    base  = _n2w(whole)
    if abs(frac - 0.5) < 0.01:
        return base + ' and one-half'
    return base


# ── Core fill function ─────────────────────────────────────────────────────────
def _replace_in_para(para, replacements):
    full_text = para.text
    if not any(k in full_text for k in replacements):
        return
    runs = para.runs
    if not runs:
        return
    combined = ''.join(r.text for r in runs)
    new_text = combined
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if new_text == combined:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def _fill_doc(doc, replacements):
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)


def _docx_to_pdf(docx_path, output_path):
    out_dir = os.path.dirname(output_path)
    result = subprocess.run(
        ['soffice', '--headless', '--convert-to', 'pdf',
         '--outdir', out_dir, docx_path],
        capture_output=True, text=True, timeout=60
    )
    base = os.path.splitext(os.path.basename(docx_path))[0]
    generated = os.path.join(out_dir, base + '.pdf')
    if os.path.exists(generated) and generated != output_path:
        os.rename(generated, output_path)
    if not os.path.exists(output_path):
        raise RuntimeError(f'PDF conversion failed: {result.stderr}')
    return output_path


# ── Main generator ─────────────────────────────────────────────────────────────
def generate_land_contract(data, output_path=None):
    template_path = os.path.abspath(os.path.join(
        os.path.dirname(__file__), '..', 'doc_templates', 'TEMPLATE_land_contract.docx'
    ))

    sale_price   = float(data['sale_price'])
    down_payment = float(data['down_payment'])
    financed     = sale_price - down_payment
    rate         = float(data['interest_rate'])
    term         = int(data['term_years'])
    pi           = float(data['pi_amount'])
    tax          = float(data['tax_monthly'])
    ins          = float(data['insurance_monthly'])
    total        = pi + tax + ins
    start_month  = data['start_month']
    start_year   = data['start_year']
    balloon      = data['balloon_date']

    replacements = {
        '[[SELLER_NAME]]':           data['seller_name'],
        '[[SELLER_ADDRESS]]':        data['seller_address'],
        '[[SELLER_SIGNATORY]]':      data['seller_signatory'],
        '[[BUYER_NAME]]':            data['buyer_name'],
        '[[PROPERTY_ADDRESS]]':      data['property_address'],
        '[[PARCEL_NO]]':             data['parcel_no'],
        '[[LEGAL_DESCRIPTION]]':     data['legal_description'],
        '[[SALE_PRICE_WORDS]]':      d2w(sale_price),
        '[[SALE_PRICE]]':            f'${sale_price:,.2f}',
        '[[DOWN_PAYMENT_WORDS]]':    d2w(down_payment),
        '[[DOWN_PAYMENT]]':          f'${down_payment:,.2f}',
        '[[FINANCED_AMOUNT_WORDS]]': d2w(financed),
        '[[FINANCED_AMOUNT]]':       f'${financed:,.2f}',
        '[[RATE_WORDS]]':            pct2w(rate),
        '[[RATE]]':                  str(int(rate)) if rate == int(rate) else str(rate),
        '[[TERM_WORDS]]':            n2w(term),
        '[[TERM_YEARS]]':            str(term),
        '[[PI_WORDS]]':              d2w(pi),
        '[[PI_AMOUNT]]':             f'${pi:,.2f}',
        '[[START_MONTH]]':           start_month,
        '[[START_YEAR]]':            start_year,
        '[[BALLOON_DATE]]':          balloon,
        '[[MONTHLY_TOTAL_WORDS]]':   d2w(total),
        '[[MONTHLY_TOTAL]]':         f'${total:,.2f}',
        'Sixty Nine and 89/100 Dollars ($69.89)': f'{d2w(tax)} (${tax:,.2f})',
        'Fifty Eight and 67/100 Dollars ($59.67)': f'{d2w(ins)} (${ins:,.2f})',
    }

    prior = data.get('prior_deed_instrument', '')
    if prior:
        replacements['Prior deed Instrument Number: ___________________________'] = \
            f'Prior deed Instrument Number: {prior}'

    if output_path is None:
        safe = data['buyer_name'].replace(' ', '_').replace('/', '_')
        output_path = os.path.join(tempfile.gettempdir(), f'LandContract_{safe}.pdf')

    with tempfile.TemporaryDirectory() as tmpdir:
        working_docx = os.path.join(tmpdir, 'land_contract_filled.docx')
        shutil.copy2(template_path, working_docx)
        doc = Document(working_docx)
        _fill_doc(doc, replacements)
        doc.save(working_docx)
        _docx_to_pdf(working_docx, output_path)

    return output_path


def generate(data):
    """Entry point called by doc_maker.py dispatcher."""
    return generate_land_contract(data)


if __name__ == '__main__':
    sample = {
        'effective_date':        '1st day of April, 2026',
        'seller_name':           'Nice City Homes LLC',
        'seller_address':        '123 Main St SW, Canton OH 44702',
        'seller_signatory':      'Michael Kell',
        'buyer_name':            'John A. Smith',
        'property_address':      '1117 Arlington Ave SW, Canton OH 44706',
        'parcel_no':             '10-12345-000',
        'legal_description':     'Lot 14 of Plat No. 5, Reedsburg Heights Addition '
                                 'to the City of Canton, Stark County, Ohio.',
        'prior_deed_instrument': '2024-00012345',
        'sale_price':            95000.00,
        'down_payment':          5000.00,
        'interest_rate':         10.0,
        'term_years':            4,
        'pi_amount':             700.00,
        'start_month':           'May',
        'start_year':            '2026',
        'balloon_date':          'April 1, 2030',
        'tax_monthly':           69.89,
        'insurance_monthly':     58.67,
        'execution_year':        '2026',
    }
    out = generate_land_contract(sample, '/tmp/LandContract_Test.pdf')
    print(f'Generated: {out}')
