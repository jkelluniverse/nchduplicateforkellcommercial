"""
Cancellation of Land Installment Contract Generator
=====================================================
Fills blank lines in CANCELLATION_OF_LAND_INSTALLMENT_CONTRACT.docx
and converts to PDF using LibreOffice.

This document uses blank underscores (____) instead of [[PLACEHOLDERS]].
Each blank is uniquely identifiable by length and surrounding context.

FORM FIELDS:
  execution_day        str    e.g. "15th"
  execution_month      str    e.g. "April"
  execution_year       str    e.g. "2026"
  vendor_name          str    seller full name or entity
  vendee_name          str    buyer full name
  vendee_address       str    buyer full address
  contract_date        str    original land contract date e.g. "April 1, 2026"
  instrument_no        str    recorder instrument number
  property_description str    short property description
  property_address     str    full property address
  parcel_no            str    PPN / parcel number
  prior_deed           str    optional prior deed reference
  seller_sig_name      str    seller signature name (line 1)
  seller_sig_name2     str    seller signature name (line 2) — optional
  buyer_sig_name       str    buyer signature name (line 1)
  buyer_sig_name2      str    buyer signature name (line 2) — optional
  notary_city          str    defaults "Canton"
  notary_day           str    blank — fill at signing
  notary_month         str    blank — fill at signing
  notary_year          str    blank — fill at signing
"""

import os, shutil, subprocess, tempfile
from docx import Document


def _replace_in_para(para, replacements):
    """Replace text in paragraph, preserving formatting of first run."""
    full_text = para.text
    if not any(k in full_text for k in replacements):
        return
    runs = para.runs
    if not runs:
        return
    combined = ''.join(r.text for r in runs)
    new_text = combined
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if new_text == combined:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def _fill_doc(doc, replacements):
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)


def _docx_to_pdf(docx_path, output_path):
    out_dir = os.path.dirname(output_path)
    subprocess.run(['soffice','--headless','--convert-to','pdf',
                    '--outdir', out_dir, docx_path],
                   capture_output=True, text=True, timeout=60)
    base = os.path.splitext(os.path.basename(docx_path))[0]
    generated = os.path.join(out_dir, base + '.pdf')
    if os.path.exists(generated) and generated != output_path:
        os.rename(generated, output_path)
    if not os.path.exists(output_path):
        raise RuntimeError('PDF conversion failed')
    return output_path


def generate(data, output_path=None):
    _here = os.path.dirname(os.path.abspath(__file__))
    for _td in [os.path.join(_here,'..','doc_templates'), os.path.join(_here,'doc_templates'), _here]:
        template_path = os.path.join(os.path.abspath(_td), 'CANCELLATION_OF_LAND_INSTALLMENT_CONTRACT.docx')
        if os.path.exists(template_path): break

    ex_day    = data.get('execution_day', '____')
    ex_month  = data.get('execution_month', '__________')
    ex_year   = data.get('execution_year', '2026')
    vendor    = data['vendor_name']
    vendee    = data['vendee_name']
    v_addr    = data['vendee_address']
    c_date    = data['contract_date']
    inst_no   = data['instrument_no']
    prop_desc = data['property_description']
    prop_addr = data['property_address']
    parcel    = data['parcel_no']
    prior     = data.get('prior_deed', '__________________________')
    sel_sig1  = data.get('seller_sig_name', '')
    sel_sig2  = data.get('seller_sig_name2', '')
    buy_sig1  = data.get('buyer_sig_name', '')
    buy_sig2  = data.get('buyer_sig_name2', '')
    n_city    = data.get('notary_city', 'Canton')
    n_day     = data.get('notary_day', '______')
    n_month   = data.get('notary_month', '______________')
    n_year    = data.get('notary_year', '_______')

    # Map each unique blank pattern to its replacement
    # Patterns identified by length and context from the original document
    replacements = {
        # Opening: "made this ____ day of __________, _____, between"
        'this ____ day of __________, _____,': f'this {ex_day} day of {ex_month}, {ex_year},',

        # Vendor (59 underscores)
        '_' * 59: vendor,

        # Vendee (55 underscores)
        '_' * 55: vendee,

        # Vendee address (43 underscores)
        '_' * 43: v_addr,

        # Contract date (19 underscores)
        '_' * 19: c_date,

        # Instrument number (30 underscores)
        '_' * 30: inst_no,

        # Property description (68 underscores)
        '_' * 68: prop_desc,

        # Property address (32 underscores)
        'Property Address: ' + '_' * 32: f'Property Address: {prop_addr}',

        # PPN (22 underscores)
        'PPN: ' + '_' * 22: f'PPN: {parcel}',

        # Prior deed reference (26 underscores)
        'Prior Deed Reference: ' + '_' * 26: f'Prior Deed Reference: {prior}',

        # Execution date in witness: "on the ______ day of _______________, _______"
        'on the ______ day of _______________, _______.':
            f'on the {ex_day} day of {ex_month}, {ex_year}.',

        # Seller signature lines (56 underscores x2)
        # First occurrence = seller line 1
        # Use context-aware replacement via sequential processing below
    }

    # Handle seller/buyer sig lines and notary separately
    # since they share the same underscore lengths
    # Do them after main replacements via a sequential pass on paragraphs

    if output_path is None:
        safe = vendee.replace(' ','_').replace('/','_')
        output_path = f'/tmp/CancellationLandContract_{safe}.pdf'

    with tempfile.TemporaryDirectory() as tmpdir:
        working = os.path.join(tmpdir, 'cancel_filled.docx')
        shutil.copy2(template_path, working)
        doc = Document(working)

        # Apply main replacements
        _fill_doc(doc, replacements)

        # Sequential pass for signature lines and notary blocks
        # Track which signature line we're on
        sig_56_count = 0
        notary_13_count = 0
        notary_6_count = 0
        notary_18_count = 0
        notary_7_count = 0

        for para in doc.paragraphs:
            t = para.text

            # Seller line 1 (56 underscores, first occurrence after SELLER:)
            if '_' * 56 in t:
                sig_56_count += 1
                if sig_56_count == 1:
                    new = t.replace('_'*56, sel_sig1, 1)
                elif sig_56_count == 2:
                    new = t.replace('_'*56, buy_sig1, 1)
                else:
                    # Notary appeared person lines (66 underscores handled above)
                    new = t
                if new != t and para.runs:
                    para.runs[0].text = new
                    for r in para.runs[1:]:
                        r.text = ''

            # Seller line 2 / buyer line 2 (31 underscores = notary sig line)
            if '_' * 31 in t:
                pass  # leave as signature line for notary

            # Notary location: "at _____________, Ohio"
            if '_' * 13 in t and 'Ohio' in t:
                notary_13_count += 1
                new = t.replace('_'*13, n_city, 1)
                if new != t and para.runs:
                    para.runs[0].text = new
                    for r in para.runs[1:]:
                        r.text = ''

            # Notary "this ______ day"
            if 'this ______ day of' in t:
                new = t.replace('this ______ day of', f'this {n_day} day of', 1)
                new = new.replace('__________________, _______', f'{n_month}, {n_year}', 1)
                if new != t and para.runs:
                    para.runs[0].text = new
                    for r in para.runs[1:]:
                        r.text = ''

        doc.save(working)
        _docx_to_pdf(working, output_path)

    return output_path


if __name__ == '__main__':
    sample = {
        'execution_day':        '7th',
        'execution_month':      'April',
        'execution_year':       '2026',
        'vendor_name':          'Nice City Homes LLC',
        'vendee_name':          'John A. Smith',
        'vendee_address':       '1117 Arlington Ave SW, Canton OH 44706',
        'contract_date':        'April 1, 2026',
        'instrument_no':        '2026-00012345',
        'property_description': '1117 Arlington Ave SW, Canton OH 44706',
        'property_address':     '1117 Arlington Ave SW, Canton OH 44706',
        'parcel_no':            '10-12345-000',
        'prior_deed':           '2024-00012345',
        'seller_sig_name':      'Michael Kell / Nice City Homes LLC',
        'buyer_sig_name':       'John A. Smith',
        'notary_city':          'Canton',
    }
    out = generate(sample, '/home/claude/CancellationLandContract_Test.pdf')
    print(f'Generated: {out}')
