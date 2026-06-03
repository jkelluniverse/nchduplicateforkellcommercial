"""
Letter of Acknowledgement Generator
======================================
Fills TEMPLATE_letter_of_acknowledgement.docx — buyer acknowledges
they are signing the Quit Claim Deed as security against default.

FORM FIELDS:
  buyer_name       str    purchaser full name
  property_address str    full property address
  seller_name      str    seller/signatory name (was hardcoded as Michael Kell)
  contract_day     str    day of original land contract e.g. "1st"
  contract_month   str    month of original land contract e.g. "April"
  notary_day       str    day of notarization — blank, fill at signing
  notary_month     str    month of notarization — blank, fill at signing
  notary_year      str    year of notarization — blank, fill at signing
"""

import os, shutil, subprocess, tempfile
from docx import Document


def _replace_in_para(para, replacements):
    full_text = para.text
    if not any(k in full_text for k in replacements):
        return
    runs = para.runs
    if not runs:
        return
    combined = ''.join(r.text for r in runs)
    new_text = combined
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if new_text == combined:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def _fill_doc(doc, replacements):
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)


def _docx_to_pdf(docx_path, output_path):
    out_dir = os.path.dirname(output_path)
    subprocess.run(['soffice','--headless','--convert-to','pdf',
                    '--outdir', out_dir, docx_path],
                   capture_output=True, text=True, timeout=60)
    base = os.path.splitext(os.path.basename(docx_path))[0]
    generated = os.path.join(out_dir, base + '.pdf')
    if os.path.exists(generated) and generated != output_path:
        os.rename(generated, output_path)
    if not os.path.exists(output_path):
        raise RuntimeError('PDF conversion failed')
    return output_path


def generate(data, output_path=None):
    _here = os.path.dirname(os.path.abspath(__file__))
    for _td in [os.path.join(_here,'..','doc_templates'), os.path.join(_here,'doc_templates'), _here]:
        template_path = os.path.join(os.path.abspath(_td), 'TEMPLATE_letter_of_acknowledgement.docx')
        if os.path.exists(template_path): break

    buyer       = data['buyer_name']
    prop_addr   = data['property_address']
    seller      = data.get('seller_name', 'Michael Kell')
    c_day       = data.get('contract_day', '____')
    c_month     = data.get('contract_month', '____________')
    n_day       = data.get('notary_day', '_____')
    n_month     = data.get('notary_month', '___________________')
    n_year      = data.get('notary_year', '______')

    replacements = {
        '[[BUYER_NAME]]':    buyer,
        '[[PROPERTY_ADDRESS]]': prop_addr,
        '[[Michael Kell]]':  seller,
        '[[day#]]':          c_day,
        '[[month   ]]':      c_month,
        # Notary date line: "this_____day of _______________________, ______"
        'this_____day of _______________________, ______':
            f'this {n_day} day of {n_month}, {n_year}',
    }

    if output_path is None:
        safe = buyer.replace(' ','_').replace('/','_')
        output_path = f'/tmp/LetterOfAcknowledgement_{safe}.pdf'

    with tempfile.TemporaryDirectory() as tmpdir:
        working = os.path.join(tmpdir, 'loa_filled.docx')
        shutil.copy2(template_path, working)
        doc = Document(working)
        _fill_doc(doc, replacements)
        doc.save(working)
        _docx_to_pdf(working, output_path)

    return output_path


if __name__ == '__main__':
    sample = {
        'buyer_name':       'John A. Smith',
        'property_address': '1117 Arlington Ave SW, Canton OH 44706',
        'seller_name':      'Michael Kell',
        'contract_day':     '1st',
        'contract_month':   'April',
    }
    out = generate(sample, '/home/claude/LetterOfAcknowledgement_Test.pdf')
    print(f'Generated: {out}')
