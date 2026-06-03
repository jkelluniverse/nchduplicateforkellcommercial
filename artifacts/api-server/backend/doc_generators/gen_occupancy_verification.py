"""
Occupancy Verification Letter Generator
=========================================
Fills the original TEMPLATE Occupancy_Verification.docx and converts to PDF.
Matches the existing NCH letter format exactly.

FORM FIELDS:
  letter_date           str    e.g. "March 23, 2026" — defaults to today
  occupant_name         str    e.g. "Joy A. Resendiz, Matthew Resendiz"
  property_address      str    full address e.g. "1815 3rd St. SE, Canton, Ohio 44707"
  parcel_no             str    e.g. "221015"
  purpose               str    dropdown selection
  purpose_other         str    only if purpose = Other
  signatory_name        str    defaults "Michael Kell"
  signatory_title       str    defaults "Property Owner"
  signatory_address     str    defaults "6521 Beverly Ave. NE, Canton, Ohio 44721"
"""

import os
import re
import shutil
import subprocess
import tempfile
from datetime import date
from docx import Document


def _replace_in_para(para, replacements):
    full_text = para.text
    if not any(k in full_text for k in replacements):
        return
    runs = para.runs
    if not runs:
        return
    combined = ''.join(r.text for r in runs)
    new_text = combined
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if new_text == combined:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def _fill_doc(doc, replacements):
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)


def _docx_to_pdf(docx_path, output_path):
    out_dir = os.path.dirname(output_path)
    subprocess.run(
        ['soffice', '--headless', '--convert-to', 'pdf',
         '--outdir', out_dir, docx_path],
        capture_output=True, text=True, timeout=60
    )
    base = os.path.splitext(os.path.basename(docx_path))[0]
    generated = os.path.join(out_dir, base + '.pdf')
    if os.path.exists(generated) and generated != output_path:
        os.rename(generated, output_path)
    if not os.path.exists(output_path):
        raise RuntimeError('PDF conversion failed')
    return output_path


def generate(data, output_path=None):
    _here = os.path.dirname(os.path.abspath(__file__))
    for _td in [os.path.join(_here,'..','doc_templates'), os.path.join(_here,'doc_templates'), _here]:
        template_path = os.path.join(os.path.abspath(_td), 'Occupancy_Verification.docx')
        if os.path.exists(template_path): break

    today = date.today().strftime('%B %d, %Y')
    letter_date    = data.get('letter_date', today)
    occupant       = data['occupant_name']
    prop_addr      = data['property_address']
    parcel         = data['parcel_no']
    purpose        = data.get('purpose', 'Utility service establishment')
    purpose_other  = data.get('purpose_other', '')
    purpose_text   = purpose_other if purpose == 'Other' and purpose_other else purpose
    sig_name       = data.get('signatory_name', 'Michael Kell')
    sig_title      = data.get('signatory_title', 'Property Owner')
    sig_addr       = data.get('signatory_address', '6521 Beverly Ave. NE, Canton, Ohio 44721')

    # The template has hardcoded values — replace them with the form data
    replacements = {
        # Date line at top
        'March 23, 2026':                    letter_date,
        # RE line
        'RE: Occupancy Verification \u2014 1815 3rd St. SE, Canton, Ohio 44707':
            f'RE: Occupancy Verification \u2014 {prop_addr}',
        # Occupant name
        'Joy A. Resendiz, Matthew Resendiz':  occupant,
        # Property address in body
        '1815 3rd St. SE, Canton, Ohio 44707': prop_addr,
        # Parcel number
        '221015':                             parcel,
        # Purpose
        'Utility service establishment only': purpose_text,
        # Signatory
        'Michael Kell':                       sig_name,
        'Property Owner':                     sig_title,
        '6521 Beverly Ave. NE, Canton, Ohio 44721': sig_addr,
    }

    if output_path is None:
        safe = occupant.replace(' ','_').replace('/','_').replace(',','')
        output_path = f'/tmp/OccupancyVerification_{safe}.pdf'

    with tempfile.TemporaryDirectory() as tmpdir:
        working = os.path.join(tmpdir, 'occ_ver_filled.docx')
        shutil.copy2(template_path, working)
        doc = Document(working)
        _fill_doc(doc, replacements)
        doc.save(working)
        _docx_to_pdf(working, output_path)

    return output_path


if __name__ == '__main__':
    sample = {
        'letter_date':       'April 7, 2026',
        'occupant_name':     'John A. Smith',
        'property_address':  '1117 Arlington Ave SW, Canton, Ohio 44706',
        'parcel_no':         '10-12345-000',
        'purpose':           'Utility service establishment',
        'signatory_name':    'Michael Kell',
        'signatory_title':   'Property Owner',
        'signatory_address': '6521 Beverly Ave. NE, Canton, Ohio 44721',
    }
    out = generate(sample, '/home/claude/OccupancyVerification_Test.pdf')
    print(f'Generated: {out}')
