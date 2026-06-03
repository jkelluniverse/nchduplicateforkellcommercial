"""
Recurring Charge Authorization Generator — NCH
================================================
Fills TEMPLATE_recurring_charge_auth.docx with tenant name,
property address, and rent amount. All payment method fields
(ACH account numbers, card details) are left blank for handwriting
at the time of signing.

FORM FIELDS:
  tenant_name       str    full legal name of tenant/purchaser
  property_address  str    full property address
  rent_amount       str    monthly rent e.g. "1,250.00"
"""

import os
import shutil
import subprocess
import tempfile
from docx import Document


def _replace_in_para(para, replacements):
    full_text = para.text
    if not any(k in full_text for k in replacements):
        return
    runs = para.runs
    if not runs:
        return
    combined = "".join(r.text for r in runs)
    new_text = combined
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if new_text == combined:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _fill_doc(doc, replacements):
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)


def generate(data):
    """
    Entry point called by doc_maker.py.
    Returns path to a temporary PDF file.
    """
    _here = os.path.dirname(os.path.abspath(__file__))
    template_path = None
    for _td in [
        os.path.join(_here, "..", "doc_templates"),
        os.path.join(_here, "doc_templates"),
        _here,
    ]:
        candidate = os.path.join(os.path.abspath(_td), "TEMPLATE_recurring_charge_auth.docx")
        if os.path.exists(candidate):
            template_path = candidate
            break

    if template_path is None:
        raise FileNotFoundError("TEMPLATE_recurring_charge_auth.docx not found")

    tenant = data.get("tenant_name", "")
    prop_addr = data.get("property_address", "")
    rent_raw = data.get("rent_amount", "")

    if rent_raw:
        clean = str(rent_raw).replace("$", "").replace(",", "").strip()
        try:
            rent_display = f"{float(clean):,.2f}"
        except ValueError:
            rent_display = clean
    else:
        rent_display = "____________"

    replacements = {
        "[[TENANT_NAME]]": tenant,
        "[[PROPERTY_ADDRESS]]": prop_addr,
        "[[RENT_AMOUNT]]": rent_display,
    }

    tmpdir = tempfile.mkdtemp()
    working = os.path.join(tmpdir, "rca_filled.docx")
    shutil.copy2(template_path, working)

    doc = Document(working)
    _fill_doc(doc, replacements)
    doc.save(working)

    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, working],
        capture_output=True,
        text=True,
        timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

    pdf_path = os.path.join(tmpdir, "rca_filled.pdf")
    if not os.path.exists(pdf_path):
        raise RuntimeError("LibreOffice did not produce rca_filled.pdf")

    return pdf_path
