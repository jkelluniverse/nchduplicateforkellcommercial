"""
Land Contract Generator — NCH as Seller (Dual Signatory)
==========================================================
Use this generator when Nice City Homes LLC is the seller.
Seller is hardcoded as NCH with two member signatories:
  Michael T. Kell (Member) and John M. Kanam (Member)

Template: backend/doc_templates/TEMPLATE_land_contract_NCH.docx

FORM FIELDS:
  buyer_name            str    purchaser full legal name
  property_address      str    full property address (also buyer's address)
  parcel_no             str    Stark County parcel number
  legal_description     str    full legal description verbatim
  prior_deed_instrument str    optional recorder instrument number
  sale_price            float  e.g. 97500.00
  down_payment          float  e.g. 2500.00
  interest_rate         float  e.g. 13.0
  term_years            int    e.g. 25
  pi_amount             float  monthly principal + interest
  start_month           str    e.g. "April"
  start_year            str    e.g. "2026"
  balloon_date          str    e.g. "April 1, 2030"
  tax_monthly           float  monthly tax escrow
  insurance_monthly     float  monthly insurance escrow
  execution_year        str    defaults to "2026"
  seller_signatory      str    defaults to "Michael T. Kell" — editable
  seller_signatory_2    str    defaults to "John M. Kanam" — editable

HARDCODED (not in form):
  Seller name:    Nice City Homes LLC, an Ohio Limited Liability Company
  Seller address: 6521 Beverly Ave NE, Canton, Ohio 44721
"""

import os, shutil, subprocess, tempfile
from docx import Document


# ── Number to words ────────────────────────────────────────────────────────────
_ones = ['','One','Two','Three','Four','Five','Six','Seven','Eight','Nine',
         'Ten','Eleven','Twelve','Thirteen','Fourteen','Fifteen','Sixteen',
         'Seventeen','Eighteen','Nineteen']
_tens = ['','','Twenty','Thirty','Forty','Fifty','Sixty','Seventy','Eighty','Ninety']

def _n2w(n):
    n = int(n)
    if n == 0:  return 'Zero'
    if n < 20:  return _ones[n]
    if n < 100: return _tens[n//10] + ('-'+_ones[n%10] if n%10 else '')
    if n < 1_000:
        return _ones[n//100]+' Hundred'+(' '+_n2w(n%100) if n%100 else '')
    if n < 1_000_000:
        return _n2w(n//1000)+' Thousand'+(' '+_n2w(n%1000) if n%1000 else '')
    return _n2w(n//1_000_000)+' Million'+(' '+_n2w(n%1_000_000) if n%1_000_000 else '')

def d2w(amount):
    dollars, cents = int(amount), round((amount - int(amount)) * 100)
    return _n2w(dollars) + (f' and {cents}/100' if cents else ' and no/100')

def n2w(n): return _n2w(int(n))

def pct2w(rate):
    whole, frac = int(rate), rate - int(rate)
    return _n2w(whole) + (' and one-half' if abs(frac-0.5) < 0.01 else '')


# ── Core fill helpers ──────────────────────────────────────────────────────────
def _replace_in_para(para, replacements):
    full_text = para.text
    if not any(k in full_text for k in replacements):
        return
    runs = para.runs
    if not runs:
        return
    combined = ''.join(r.text for r in runs)
    new_text = combined
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if new_text == combined:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def _fill_doc(doc, replacements):
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)


def _docx_to_pdf(docx_path, output_path):
    out_dir = os.path.dirname(output_path)
    subprocess.run(
        ['soffice', '--headless', '--convert-to', 'pdf',
         '--outdir', out_dir, docx_path],
        capture_output=True, text=True, timeout=60
    )
    base = os.path.splitext(os.path.basename(docx_path))[0]
    generated = os.path.join(out_dir, base + '.pdf')
    if os.path.exists(generated) and generated != output_path:
        os.rename(generated, output_path)
    if not os.path.exists(output_path):
        raise RuntimeError('PDF conversion failed')
    return output_path


# ── Main generator ─────────────────────────────────────────────────────────────
def generate_land_contract_nch(data, output_path=None):
    """
    Generate a land contract PDF with Nice City Homes LLC as the seller.
    Seller name and address are hardcoded. Two member signatories.
    """
    _here = os.path.dirname(os.path.abspath(__file__))
    for _td in [os.path.join(_here, '..', 'doc_templates'),
                os.path.join(_here, 'doc_templates'), _here]:
        template_path = os.path.join(os.path.abspath(_td),
                                     'TEMPLATE_land_contract_NCH.docx')
        if os.path.exists(template_path):
            break

    buyer       = data['buyer_name']
    prop_addr   = data['property_address']
    parcel      = data['parcel_no']
    legal       = data['legal_description']
    prior       = data.get('prior_deed_instrument', '')
    sale_price  = float(data['sale_price'])
    down        = float(data['down_payment'])
    financed    = sale_price - down
    rate        = float(data['interest_rate'])
    term        = int(data['term_years'])
    pi          = float(data['pi_amount'])
    start_month = data['start_month']
    start_year  = data['start_year']
    balloon     = data['balloon_date']
    tax         = float(data['tax_monthly'])
    ins         = float(data['insurance_monthly'])
    total       = pi + tax + ins
    exec_year   = data.get('execution_year', start_year)
    sig1        = data.get('seller_signatory', 'Michael T. Kell')
    sig2        = data.get('seller_signatory_2', 'John M. Kanam')

    if output_path is None:
        safe = buyer.replace(' ', '_').replace('/', '_')
        output_path = f'/tmp/LandContract_NCH_{safe}.pdf'

    replacements = {
        '[[BUYER_NAME]]':            buyer,
        '[[PROPERTY_ADDRESS]]':      prop_addr,
        '[[PARCEL_NO]]':             parcel,
        '[[LEGAL_DESCRIPTION]]':     legal,
        '[[SALE_PRICE_WORDS]]':      d2w(sale_price),
        '[[SALE_PRICE]]':            f'${sale_price:,.2f}',
        '[[DOWN_PAYMENT_WORDS]]':    d2w(down),
        '[[DOWN_PAYMENT]]':          f'${down:,.2f}',
        '[[FINANCED_AMOUNT_WORDS]]': d2w(financed),
        '[[FINANCED_AMOUNT]]':       f'${financed:,.2f}',
        '[[RATE_WORDS]]':            pct2w(rate),
        '[[RATE]]':                  str(int(rate)) if rate == int(rate) else str(rate),
        '[[TERM_WORDS]]':            n2w(term),
        '[[TERM_YEARS]]':            str(term),
        '[[PI_WORDS]]':              d2w(pi),
        '[[PI_AMOUNT]]':             f'${pi:,.2f}',
        '[[START_MONTH]]':           start_month,
        '[[START_YEAR]]':            start_year,
        '[[BALLOON_DATE]]':          balloon,
        '[[MONTHLY_TOTAL_WORDS]]':   d2w(total),
        '[[MONTHLY_TOTAL]]':         f'${total:,.2f}',
        '[[SELLER_SIGNATORY]]':      sig1,
        '[[SELLER_SIGNATORY_2]]':    sig2,
        # Hardcoded tax/insurance placeholders in Section 3
        'Sixty Nine and 89/100 Dollars ($69.89)': f'{d2w(tax)} (${tax:,.2f})',
        'Fifty Eight and 67/100 Dollars ($59.67)': f'{d2w(ins)} (${ins:,.2f})',
    }

    if prior:
        replacements['Prior deed Instrument Number: ___________________________'] = \
            f'Prior deed Instrument Number: {prior}'

    with tempfile.TemporaryDirectory() as tmpdir:
        working = os.path.join(tmpdir, 'lc_nch_filled.docx')
        shutil.copy2(template_path, working)
        doc = Document(working)
        _fill_doc(doc, replacements)
        doc.save(working)
        _docx_to_pdf(working, output_path)

    return output_path


def generate(data):
    """Entry point called by doc_maker.py."""
    import tempfile
    safe = data.get('buyer_name', 'buyer').replace(' ', '_').replace('/', '_')
    output_path = os.path.join(tempfile.gettempdir(), f'LandContract_NCH_{safe}.pdf')
    return generate_land_contract_nch(data, output_path)
