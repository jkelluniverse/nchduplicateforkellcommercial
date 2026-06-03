"""
Acknowledgement and Hold Harmless Agreement Generator
=======================================================
Fills TEMPLATE_acknowledgement_harmless.docx — buyer acknowledges
AS-IS purchase risks and releases NCH from liability.

FORM FIELDS:
  property_address  str    full property address
  buyer_name        str    purchaser full name
  sign_date         str    date of signing — defaults to today
"""

import os, shutil, subprocess, tempfile
from datetime import date
from docx import Document


def _replace_in_para(para, replacements):
    full_text = para.text
    if not any(k in full_text for k in replacements):
        return
    runs = para.runs
    if not runs:
        return
    combined = ''.join(r.text for r in runs)
    new_text = combined
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if new_text == combined:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def _fill_doc(doc, replacements):
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)


def _docx_to_pdf(docx_path, output_path):
    out_dir = os.path.dirname(output_path)
    subprocess.run(['soffice','--headless','--convert-to','pdf',
                    '--outdir', out_dir, docx_path],
                   capture_output=True, text=True, timeout=60)
    base = os.path.splitext(os.path.basename(docx_path))[0]
    generated = os.path.join(out_dir, base + '.pdf')
    if os.path.exists(generated) and generated != output_path:
        os.rename(generated, output_path)
    if not os.path.exists(output_path):
        raise RuntimeError('PDF conversion failed')
    return output_path


def generate(data, output_path=None):
    _here = os.path.dirname(os.path.abspath(__file__))
    for _td in [os.path.join(_here,'..','doc_templates'), os.path.join(_here,'doc_templates'), _here]:
        template_path = os.path.join(os.path.abspath(_td), 'TEMPLATE_acknowledgement_harmless.docx')
        if os.path.exists(template_path): break

    prop_addr  = data['property_address']
    buyer      = data['buyer_name']
    sign_date  = data.get('sign_date', date.today().strftime('%B %d, %Y'))

    replacements = {
        '[[PROPERTY_ADDRESS]]': prop_addr,
        '[[BUYER_NAME]]':       buyer,
        # Date line: "Date:____________________"
        'Date:____________________': f'Date: {sign_date}',
    }

    if output_path is None:
        safe = buyer.replace(' ','_').replace('/','_')
        output_path = f'/tmp/HoldHarmless_{safe}.pdf'

    with tempfile.TemporaryDirectory() as tmpdir:
        working = os.path.join(tmpdir, 'hh_filled.docx')
        shutil.copy2(template_path, working)
        doc = Document(working)
        _fill_doc(doc, replacements)
        doc.save(working)
        _docx_to_pdf(working, output_path)

    return output_path


if __name__ == '__main__':
    sample = {
        'property_address': '1117 Arlington Ave SW, Canton OH 44706',
        'buyer_name':       'John A. Smith',
        'sign_date':        'April 7, 2026',
    }
    out = generate(sample, '/home/claude/HoldHarmless_Test.pdf')
    print(f'Generated: {out}')
