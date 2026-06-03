"""
Quit Claim Deed Generator — 2026 Format (Template Fill)
=========================================================
Fills [[PLACEHOLDERS]] in TEMPLATE_quit_claim.docx and converts to PDF
using LibreOffice. Output is pixel-perfect match to the source document.

Template file: backend/doc_templates/TEMPLATE_quit_claim.docx
LibreOffice must be installed on the server.

Supports two grantor types:

TYPE 1 — INDIVIDUAL:
  grantor_type         = "individual"
  grantor_name         = "John A. Smith"
  grantor_entity_type  = ""
  grantor_signatory    = "John A. Smith"
  grantor_member_title = ""

TYPE 2 — COMPANY/LLC:
  grantor_type         = "company"
  grantor_name         = "NICE CITY HOMES LLC"
  grantor_entity_type  = "an Ohio Limited Liability Company"
  grantor_signatory    = "Michael T. Kell"
  grantor_member_title = "Sole Member"

COMMON FIELDS:
  grantee_name          str
  property_address      str
  parcel_no             str
  legal_description     str
  prior_deed_reference  str  optional
  execution_year        str  e.g. "2026"
"""

import os
import shutil
import subprocess
import tempfile
from docx import Document


_year_words = {
    '2024': 'Two Thousand Twenty Four',
    '2025': 'Two Thousand Twenty Five',
    '2026': 'Two Thousand Twenty Six',
    '2027': 'Two Thousand Twenty Seven',
    '2028': 'Two Thousand Twenty Eight',
    '2029': 'Two Thousand Twenty Nine',
    '2030': 'Two Thousand Thirty',
}


def year_to_words(year_str):
    return _year_words.get(str(year_str), f'Two Thousand {str(year_str)[-2:]}')


# ── Core fill helpers ──────────────────────────────────────────────────────────
def _replace_in_para(para, replacements):
    full_text = para.text
    if not any(k in full_text for k in replacements):
        return
    runs = para.runs
    if not runs:
        return
    combined = ''.join(r.text for r in runs)
    new_text = combined
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if new_text == combined:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def _fill_doc(doc, replacements):
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)


def _docx_to_pdf(docx_path, output_path):
    out_dir = os.path.dirname(output_path)
    subprocess.run(
        ['soffice', '--headless', '--convert-to', 'pdf',
         '--outdir', out_dir, docx_path],
        capture_output=True, text=True, timeout=60
    )
    base = os.path.splitext(os.path.basename(docx_path))[0]
    generated = os.path.join(out_dir, base + '.pdf')
    if os.path.exists(generated) and generated != output_path:
        os.rename(generated, output_path)
    if not os.path.exists(output_path):
        raise RuntimeError('PDF conversion failed')
    return output_path


def generate_quit_claim_deed(data, output_path=None):
    _here = os.path.dirname(os.path.abspath(__file__))
    template_path = None
    for _td in [os.path.join(_here, '..', 'doc_templates'),
                os.path.join(_here, 'doc_templates'), _here]:
        candidate = os.path.join(os.path.abspath(_td),
                                 'TEMPLATE_quit_claim.docx')
        if os.path.exists(candidate):
            template_path = candidate
            break
    if template_path is None:
        raise FileNotFoundError(
            'TEMPLATE_quit_claim.docx not found in doc_templates'
        )

    grantor_type   = (data.get('grantor_type') or 'individual').lower()
    grantor_name   = data['grantor_name']
    grantor_entity = data.get('grantor_entity_type', '') or ''
    grantor_sig    = data.get('grantor_signatory') or grantor_name
    member_title   = data.get('grantor_member_title') or 'Sole Member'
    grantee        = data['grantee_name']
    prop_addr      = data['property_address']
    parcel         = data['parcel_no']
    legal          = data['legal_description']
    prior_deed     = (data.get('prior_deed_reference')
                      or data.get('prior_deed_instrument')
                      or '___________________________')
    exec_year      = str(data.get('execution_year', '2026'))

    if grantor_type == 'individual' or not grantor_entity:
        grantor_entity = ''
        member_title   = ''

    # Collapse the ", [[GRANTOR_ENTITY_TYPE]]" into nothing for individuals
    if grantor_entity:
        grantor_name_full = '[[GRANTOR_NAME]], [[GRANTOR_ENTITY_TYPE]]'
        grantor_name_repl = f'{grantor_name}, {grantor_entity}'
    else:
        grantor_name_full = '[[GRANTOR_NAME]], [[GRANTOR_ENTITY_TYPE]]'
        grantor_name_repl = grantor_name

    replacements = {
        grantor_name_full:          grantor_name_repl,
        '[[GRANTOR_NAME]]':         grantor_name,
        '[[GRANTOR_ENTITY_TYPE]]':  grantor_entity,
        '[[GRANTEE_NAME]]':         grantee,
        '[[PROPERTY_ADDRESS]]':     prop_addr,
        '[[PARCEL_NO]]':            parcel,
        '[[GRANTOR_SIGNATORY]]':    grantor_sig,
        '[[GRANTOR_MEMBER_TITLE]]': member_title,
        '[[LEGAL_DESCRIPTION]]':    legal,
        '[[PRIOR_DEED_REFERENCE]]': prior_deed,
        '[[EXECUTION_YEAR_WORDS]]': year_to_words(exec_year),
        '[[EXECUTION_YEAR]]':       exec_year,
    }

    if output_path is None:
        safe = grantee.replace(' ', '_').replace('/', '_')
        output_path = os.path.join(tempfile.gettempdir(),
                                   f'QuitClaimDeed_{safe}.pdf')

    with tempfile.TemporaryDirectory() as tmpdir:
        working = os.path.join(tmpdir, 'quit_claim_filled.docx')
        shutil.copy2(template_path, working)
        doc = Document(working)
        _fill_doc(doc, replacements)
        doc.save(working)
        _docx_to_pdf(working, output_path)

    return output_path


def generate(data):
    """Entry point called by doc_maker.py dispatcher."""
    return generate_quit_claim_deed(data)


if __name__ == '__main__':
    sample = {
        'grantor_type':         'company',
        'grantor_name':         'ALMNINI TRADES, LLC',
        'grantor_entity_type':  'an Ohio Limited Liability Company',
        'grantor_signatory':    'TAIM ALLAH ALMNINI',
        'grantor_member_title': 'Sole Member',
        'grantee_name':         'Jacob Kell',
        'property_address':     '1946 Otto Pl. NE Canton, Ohio 44704',
        'parcel_no':            '210306',
        'legal_description':    ('Known as and being eighty (80) feet off the entire '
                                 'North end of Lot No. Six Thousand Thirty-three (#6033) '
                                 'in the George Gross Allotment, as shown in Plat Volume '
                                 '3, Page 70 of the Stark County Ohio Records.'),
        'prior_deed_reference': '202602090004704',
        'execution_year':       '2026',
    }
    out = generate(sample)
    print(f'Generated: {out}')
